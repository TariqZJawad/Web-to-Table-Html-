import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class CVInfo:
    """كلاس مخصص لتجميع واستقبال بيانات المستخدم"""
    def __init__(self):
        self.name = ""
        self.title = ""
        self.email = ""
        self.phone = ""
        self.summary = ""
        self.experience = []  # قائمة لتخزين الخبرات
        self.education = []   # قائمة لتخزين التعليم

    def collect_data(self):
        """دالة تفاعلية لطلب البيانات من المستخدم عبر الترمينال"""
        print("=== 📝 مرحباً بك في منشئ السير الذاتية الاحترافي ===")
        self.name = input("الاسم الكامل (بالإنجليزي يفضل): ")
        self.title = input("المسمى الوظيفي (مثلاً: Python Developer): ")
        self.email = input("البريد الإلكتروني: ")
        self.phone = input("رقم الهاتف: ")
        self.summary = input("نبذة مختصرة عنك (Summary): ")
        
        # جمع الخبرات المهنية
        print("\n--- 💼 أقسام الخبرة المهنية (أدخل 'q' لإنهاء القسم) ---")
        while True:
            job_title = input("المسمى الوظيفي للخبرة: ")
            if job_title.lower() == 'q': break
            company = input("الشركة / المؤسسة: ")
            years = input("المدة الزمنية (مثلاً: 2024 - 2026): ")
            self.experience.append({"title": job_title, "company": company, "years": years})
            print("  تمت إضافة الخبرة بنجاح. (أدخل الخبرة التالية أو 'q' للخروج)")

        # جمع التعليم
        print("\n--- 🎓 أقسام التعليم والشهادات (أدخل 'q' لإنهاء القسم) ---")
        while True:
            degree = input("الشهادة / التخصص: ")
            if degree.lower() == 'q': break
            university = input("الجامعة / المعهد: ")
            year = input("سنة التخرج: ")
            self.education.append({"degree": degree, "university": university, "year": year})
            print("  تمت إضافة التعليم بنجاح.")


class CVGenerator:
    """كلاس مسؤول عن هندسة وتنسيق ملف الوورد بناءً على البيانات"""
    def __init__(self, cv_info: CVInfo):
        self.cv_info = cv_info
        self.doc = Document()
        
    def _set_font(self, run, font_name="Arial", size_pt=11, bold=False):
        """دالة داخلية مساعدة لتنسيق الخطوط بسرعة منعاً للتكرار (DRY)"""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold

    def build_word_file(self, filename="My_Professional_CV.docx"):
        """بناء الهيكل الإنشائي للملف"""
        # 1. العنوان الرئيسي (الاسم والمسمى)
        header_p = self.doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        name_run = header_p.add_run(f"{self.cv_info.name}\n")
        self._set_font(name_run, size_pt=24, bold=True)
        
        title_run = header_p.add_run(f"{self.cv_info.title}\n")
        self._set_font(title_run, size_pt=14, bold=False)
        
        contact_run = header_p.add_run(f"📧 {self.cv_info.email}  |  📱 {self.cv_info.phone}\n")
        self._set_font(contact_run, size_pt=10)
        
        self.doc.add_paragraph("-" * 60) # خط فاصل طفيف

        # 2. قسم النبذة الشخصية (Summary)
        self.doc.add_heading("Professional Summary", level=1)
        summary_p = self.doc.add_paragraph(self.cv_info.summary)
        
        # 3. قسم الخبرات (Work Experience)
        self.doc.add_heading("Work Experience", level=1)
        for exp in self.cv_info.experience:
            p = self.doc.add_paragraph()
            lead_run = p.add_run(f"• {exp['title']} ")
            self._set_font(lead_run, bold=True)
            
            info_run = p.add_run(f"at {exp['company']} ({exp['years']})")
            self._set_font(info_run)

        # 4. قسم التعليم (Education)
        self.doc.add_heading("Education", level=1)
        for edu in self.cv_info.education:
            p = self.doc.add_paragraph()
            edu_run = p.add_run(f"• {edu['degree']} ")
            self._set_font(edu_run, bold=True)
            
            uni_run = p.add_run(f"- {edu['university']} ({edu['year']})")
            self._set_font(uni_run)

        # 5. حفظ الملف النهائي
        self.doc.save(filename)
        print(f"\n🎉 رائع! تم إنشاء ملف السيرة الذاتية بنجاح باسم: {os.path.abspath(filename)}")


# --- نقطة انطلاق البرنامج الفعلي ---
if __name__ == "__main__":
    # إنشاء كائن البيانات وتشغيل واجهة استقبال المدخلات
    user_cv = CVInfo()
    user_cv.collect_data()
    
    # تمرير البيانات لكائن المولد لتصدير ملف الوورد
    generator = CVGenerator(user_cv)
    generator.build_word_file()
